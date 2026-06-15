from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas


DOCX_OUT = Path("Saatvik_Raghuvanshi_Resume_ATS.docx")
PDF_OUT = Path("Saatvik_Raghuvanshi_Resume_ATS.pdf")


RESUME = {
    "name": "Saatvik Raghuvanshi",
    "headline": "B.Tech CSE Student | Full-Stack Developer | Cloud Engineering Aspirant",
    "contact": [
        ("Jaipur, Rajasthan, India", None),
        ("+91 9301661150", None),
        ("raghuvanshisaatvik@gmail.com", "mailto:raghuvanshisaatvik@gmail.com"),
        ("linkedin.com/in/saatvik-raghuvanshi", "https://www.linkedin.com/in/saatvik-raghuvanshi"),
        ("github.com/saatvikraghuvanshi-lab", "https://github.com/saatvikraghuvanshi-lab"),
    ],
    "sections": [
        {
            "heading": "SUMMARY",
            "paragraphs": [
                "Computer Science undergraduate at Manipal University Jaipur with hands-on experience building deployed full-stack web applications, AI-assisted workflows, Supabase/PostgreSQL backends, and frontend-focused product interfaces. Comfortable working across React, Next.js, TypeScript, Tailwind CSS, authentication, database design, cloud deployment, and practical UI/UX through hackathons, ideathons, freelance work, and project-based learning.",
            ],
        },
        {
            "heading": "TECHNICAL SKILLS",
            "pairs": [
                ("Languages", "C, Python, TypeScript, JavaScript, HTML, CSS"),
                ("Frontend", "React.js, Next.js, Tailwind CSS, shadcn/ui, Radix UI, Redux, responsive UI/UX"),
                ("Backend and Database", "Node.js, PostgreSQL, Supabase, Convex, REST APIs, authentication, Row Level Security"),
                ("AI, Cloud and Tools", "Gemini API, Inngest, Firebase, Google Cloud, Vercel, Netlify, Git, GitHub, ngrok"),
                ("Other Technical Skills", "AutoCAD, drone building, line follower robots, GIS map integration, robotics fundamentals"),
            ],
        },
        {
            "heading": "EDUCATION",
            "roles": [
                {
                    "title": "Manipal University Jaipur - B.Tech in Computer Science and Engineering",
                    "meta": "Expected 2029 | Current GPA: 7.0",
                    "bullets": [
                        "Completed first year while building practical projects across web development, APIs, databases, AI workflows, UI/UX, and cloud-related tools.",
                    ],
                },
                {
                    "title": "Army Public School - Class 10: 86% | Resonance - Class 12: 82%",
                    "bullets": [
                        "Built early technical exposure through robotics, drone building, line follower systems, and technical workshops.",
                    ],
                },
            ],
        },
        {
            "heading": "PROJECTS",
            "roles": [
                {
                    "title": "ShockProof - AI Smart Meter Tariff Guard",
                    "meta": "Live: https://shockproof.vercel.app | GitHub: github.com/saatvikraghuvanshi-lab/ShockProof",
                    "bullets": [
                        "Built an AI-powered electricity meter assistant for Indian households using Next.js, React, TypeScript, Tailwind CSS, Supabase Auth/Database/Storage/Realtime, Gemini OCR, and Vercel.",
                        "Implemented secure meter photo/video uploads, reading history, manual correction fallback, kWh extraction, month-end usage projection, tariff slab distance, estimated bill, bill-risk level, and Gemini-generated household advice.",
                    ],
                },
                {
                    "title": "S2C - AI Sketch-to-Design Learning Project",
                    "meta": "Live: https://ai-saa-s-sketch-to-design.vercel.app | GitHub: github.com/saatvikraghuvanshi-lab/Ai-SaaS-Sketch-To-Design",
                    "bullets": [
                        "Built a full-stack AI SaaS-style design tool with Google auth, project dashboard, infinite canvas editor, Redux canvas state, Convex database/auth/storage, Inngest autosave, Gemini API routes, PNG exports, and JSON exports.",
                    ],
                },
                {
                    "title": "VibeBatch - Social Personality Web Application",
                    "meta": "Live: https://www.vibebatch.net | GitHub: github.com/saatvikraghuvanshi-lab/VibeBatch",
                    "bullets": [
                        "Architected Supabase/PostgreSQL data flows for user profiles, friend relationships, anonymous trait voting, messages, invite sharing, premium identity cards, Gemini help support, and Row Level Security-based privacy.",
                    ],
                },
                {
                    "title": "ResilienceOS - Disaster Management Web Platform",
                    "meta": "Live: https://resilienceos.vercel.app | GitHub: github.com/saatvikraghuvanshi-lab/RESILIENCEOS",
                    "bullets": [
                        "Built role-based disaster-response workflows for admin command, responder tasks, civilian SOS reporting, map simulation, readiness training, strategic forecasting, and report generation; recognized as a 4th place Startup Forge Ideathon project.",
                    ],
                },
                {
                    "title": "JanSahayak - AI Public Assistance Web Application",
                    "meta": "Live: https://jansahayak1.netlify.app",
                    "bullets": [
                        "Developed a civic-tech prototype for a college Vibeathon covering login, scheme discovery, filters, user dashboard, and AI assistant flow; placed in the top 5.",
                    ],
                },
            ],
        },
        {
            "heading": "EXPERIENCE",
            "roles": [
                {
                    "title": "Freelance Full-Stack Developer - VibeBatch",
                    "bullets": [
                        "Delivered a live freelance web product with frontend implementation, Supabase-backed workflows, authentication, database design, deployment readiness, and product UI refinement.",
                    ],
                },
                {
                    "title": "Cloud Computing Training - Training Phase",
                    "bullets": [
                        "Building foundational cloud computing knowledge aligned with cloud engineering internships, deployment workflows, backend services, and infrastructure fundamentals.",
                    ],
                },
            ],
        },
        {
            "heading": "ACHIEVEMENTS AND CERTIFICATIONS",
            "bullets": [
                "INDIA.RUNS 2026 Ideathon: Built ShockProof for Challenge 3, Improve Everyday Life with AI, as part of team Noble Dawn.",
                "Startup Forge Ideathon: 4th Place Finisher, GCEC Global Foundation, for ResilienceOS during a 48-hour build.",
                "Technical workshops and hackathons: Rewind & Recode National Hackathon, Robotics Workshop at Techfest IIT Bombay, WRC Quadcopter Challenge, Fastest Line Follower Challenge, Machine Learning Workshop, and AI/Prompt Engineering career session.",
            ],
        },
    ],
}


def set_run_font(run, name="Arial", size=None, bold=None, color="000000"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_link(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    r_pr.append(color)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def configure_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.38)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(9.6)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_after = Pt(1.35)
    normal.paragraph_format.line_spacing = 1.0

    for name in ("Resume Heading", "Resume Role", "Resume Bullet", "Resume Pair"):
        if name in styles:
            continue
        styles.add_style(name, 1)

    heading = styles["Resume Heading"]
    heading.base_style = normal
    heading.font.name = "Arial"
    heading._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    heading.font.size = Pt(10.8)
    heading.font.bold = True
    heading.paragraph_format.space_before = Pt(4.7)
    heading.paragraph_format.space_after = Pt(1.35)
    heading.paragraph_format.line_spacing = 1.0

    role = styles["Resume Role"]
    role.base_style = normal
    role.font.name = "Arial"
    role._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    role.font.size = Pt(9.55)
    role.paragraph_format.space_before = Pt(1.6)
    role.paragraph_format.space_after = Pt(0.55)
    role.paragraph_format.line_spacing = 1.0

    bullet = styles["Resume Bullet"]
    bullet.base_style = normal
    bullet.font.name = "Arial"
    bullet._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    bullet.font.size = Pt(8.95)
    bullet.paragraph_format.left_indent = Inches(0.12)
    bullet.paragraph_format.first_line_indent = Inches(-0.12)
    bullet.paragraph_format.space_after = Pt(0.65)
    bullet.paragraph_format.line_spacing = 0.98

    pair = styles["Resume Pair"]
    pair.base_style = normal
    pair.font.name = "Arial"
    pair._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    pair.font.size = Pt(8.95)
    pair.paragraph_format.space_after = Pt(0.75)
    pair.paragraph_format.line_spacing = 0.98


def add_rule(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_heading(doc, text):
    p = doc.add_paragraph(style="Resume Heading")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    add_rule(p)


def add_contact_line(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    for index, (text, url) in enumerate(RESUME["contact"]):
        if index:
            r = p.add_run(" | ")
            set_run_font(r, size=8.35)
        if url:
            set_link(p, text, url)
        else:
            r = p.add_run(text)
            set_run_font(r, size=8.35)


def add_role(doc, role):
    p = doc.add_paragraph(style="Resume Role")
    p.paragraph_format.keep_with_next = True
    title = p.add_run(role["title"])
    set_run_font(title, size=8.75, bold=True)
    if role.get("meta"):
        meta = p.add_run(f" | {role['meta']}")
        set_run_font(meta, size=8.25)
    for bullet in role.get("bullets", []):
        b = doc.add_paragraph(style="Resume Bullet")
        run = b.add_run("- " + bullet)
        set_run_font(run, size=8.25)


def build_docx():
    doc = Document()
    configure_doc(doc)

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.paragraph_format.space_after = Pt(0)
    r = name.add_run(RESUME["name"].upper())
    set_run_font(r, size=16.4, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(1)
    r = title.add_run(RESUME["headline"])
    set_run_font(r, size=8.8)
    add_contact_line(doc)

    for section in RESUME["sections"]:
        add_heading(doc, section["heading"])
        for paragraph in section.get("paragraphs", []):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1.35)
            run = p.add_run(paragraph)
            set_run_font(run, size=8.5)
        for label, value in section.get("pairs", []):
            p = doc.add_paragraph(style="Resume Pair")
            label_run = p.add_run(f"{label}: ")
            set_run_font(label_run, size=8.3, bold=True)
            value_run = p.add_run(value)
            set_run_font(value_run, size=8.3)
        for role in section.get("roles", []):
            add_role(doc, role)
        for bullet in section.get("bullets", []):
            p = doc.add_paragraph(style="Resume Bullet")
            run = p.add_run("- " + bullet)
            set_run_font(run, size=8.25)

    doc.save(DOCX_OUT)


def wrap_text(text, font, size, width):
    words = text.split()
    lines = []
    current = ""
    for word in words:
        test = word if not current else f"{current} {word}"
        if stringWidth(test, font, size) <= width:
            current = test
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines or [""]


def build_pdf():
    page_w, page_h = letter
    margin_x = 0.45 * inch
    margin_top = 0.34 * inch
    margin_bottom = 0.30 * inch
    width = page_w - 2 * margin_x
    y = page_h - margin_top
    c = canvas.Canvas(str(PDF_OUT), pagesize=letter)

    def draw_lines(text, font="Helvetica", size=9.25, leading=10.35, indent=0, after=0.85):
        nonlocal y
        lines = wrap_text(text, font, size, width - indent)
        for line in lines:
            if y < margin_bottom + leading:
                raise RuntimeError("Resume exceeded one PDF page")
            c.setFont(font, size)
            c.drawString(margin_x + indent, y, line)
            y -= leading
        y -= after

    c.setTitle("Saatvik Raghuvanshi Resume")
    c.setAuthor("Saatvik Raghuvanshi")
    c.setFont("Helvetica-Bold", 18.4)
    name_w = stringWidth(RESUME["name"].upper(), "Helvetica-Bold", 18.4)
    c.drawString((page_w - name_w) / 2, y, RESUME["name"].upper())
    y -= 15.5

    c.setFont("Helvetica", 9.85)
    head_w = stringWidth(RESUME["headline"], "Helvetica", 9.85)
    c.drawString((page_w - head_w) / 2, y, RESUME["headline"])
    y -= 12.4

    contact = " | ".join(item[0] for item in RESUME["contact"])
    c.setFont("Helvetica", 9.0)
    for line in wrap_text(contact, "Helvetica", 9.0, width):
        line_w = stringWidth(line, "Helvetica", 9.0)
        c.drawString((page_w - line_w) / 2, y, line)
        y -= 9.85
    y -= 2

    for section in RESUME["sections"]:
        y -= 2.2
        c.setFont("Helvetica-Bold", 10.8)
        c.drawString(margin_x, y, section["heading"])
        c.line(margin_x, y - 1.8, page_w - margin_x, y - 1.8)
        y -= 10.75

        for paragraph in section.get("paragraphs", []):
            draw_lines(paragraph, "Helvetica", 9.35, 10.55, after=0.95)
        for label, value in section.get("pairs", []):
            draw_lines(f"{label}: {value}", "Helvetica", 9.15, 10.25, after=0.3)
        for role in section.get("roles", []):
            line = role["title"]
            if role.get("meta"):
                line += f" | {role['meta']}"
            draw_lines(line, "Helvetica-Bold", 9.35, 10.45, after=0.2)
            for bullet in role.get("bullets", []):
                draw_lines("- " + bullet, "Helvetica", 9.05, 10.2, indent=8, after=0.3)
        for bullet in section.get("bullets", []):
            draw_lines("- " + bullet, "Helvetica", 9.05, 10.2, indent=8, after=0.3)

    c.showPage()
    c.save()


def main():
    build_docx()
    build_pdf()
    print(f"Wrote {DOCX_OUT}")
    print(f"Wrote {PDF_OUT}")


if __name__ == "__main__":
    main()
